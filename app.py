import streamlit as st
import os
from dotenv import load_dotenv
from PyPDF2 import PdfReader

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI

from langchain_core.runnables import RunnablePassthrough, RunnableMap
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.documents import Document

import pandas as pd
from docx import Document as DocxDocument

from langchain_community.retrievers import BM25Retriever
from langchain_classic.retrievers.ensemble import EnsembleRetriever

from langchain_pinecone import PineconeVectorStore
from pinecone import Pinecone

# LOAD ENV VARIABLES
load_dotenv()
os.environ["GOOGLE_API_KEY"] = os.getenv("GOOGLE_API_KEY")


# DOCUMENT EXTRACTION FUNCTIONS
def extract_docx_documents(docx_files):
    documents = []

    for file in docx_files:
        doc = DocxDocument(file)
        full_text = []

        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)

        content = "\n".join(full_text)

        if content:
            documents.append(
                Document(
                    page_content=content,
                    metadata={
                        "source": file.name,
                        "page": 1
                    }
                )
            )
    return documents


def extract_csv_documents(csv_files):
    documents = []

    for file in csv_files:
        df = pd.read_csv(file)

        for index, row in df.iterrows():
            row_text = ", ".join(
                f"{col}: {row[col]}" for col in df.columns
            )

            documents.append(
                Document(
                    page_content=row_text,
                    metadata={
                        "source": file.name,
                        "page": index + 1
                    }
                )
            )
    return documents


def extract_pdf_documents(pdf_files):
    documents = []

    for pdf in pdf_files:
        reader = PdfReader(pdf)

        for page_number, page in enumerate(reader.pages):
            content = page.extract_text()

            if content:
                documents.append(
                    Document(
                        page_content=content,
                        metadata={
                            "source": pdf.name,
                            "page": page_number + 1
                        }
                    )
                )
    return documents


# CHUNKING
def split_documents(documents):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=2000,
        chunk_overlap=100
    )
    return splitter.split_documents(documents)


# VECTOR STORE
def create_vector_store(chunks):

    embeddings = GoogleGenerativeAIEmbeddings(
        model="gemini-embedding-001"
    )

    pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
    index = pc.Index(os.getenv("PINECONE_INDEX_NAME"))

    vector_store = PineconeVectorStore(
        index=index,
        embedding=embeddings
    )

    vector_store.add_documents(chunks)


# LOAD RAG CHAIN (WITH MEMORY)
def load_rag_chain(selected_model):

    embeddings = GoogleGenerativeAIEmbeddings(
        model="gemini-embedding-001"
    )

    pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
    index = pc.Index(os.getenv("PINECONE_INDEX_NAME"))

    db = PineconeVectorStore(
        index=index,
        embedding=embeddings
    )

    # Semantic Retriever
    vector_retriever = db.as_retriever(search_kwargs={"k": 4})

    # BM25 Retriever
    docs = list(db.docstore._dict.values())
    bm25_retriever = BM25Retriever.from_documents(docs)
    bm25_retriever.k = 4

    # Hybrid Retriever
    hybrid_retriever = EnsembleRetriever(
        retrievers=[vector_retriever, bm25_retriever],
        weights=[0.6, 0.4]
    )

    # LLM
    llm = ChatGoogleGenerativeAI(
        model=selected_model,
        temperature=0.3,
        streaming=True
    )

    # Prompt with Chat History
    prompt = ChatPromptTemplate.from_template("""
You are an AI assistant answering strictly based on provided context.

Use previous conversation history if relevant.

If the answer is not found in the context, respond with:
"Answer not available in the provided documents."

Chat History:
{chat_history}

Context:
{context}

Question:
{question}
""")

    def format_docs(docs):
        return "\n\n".join(doc.page_content for doc in docs)

    def format_chat_history(history):
        formatted = ""
        for chat in history:
            formatted += f"User: {chat['question']}\n"
            formatted += f"Assistant: {chat['answer']}\n\n"
        return formatted

    retrieval_chain = RunnableMap({
        "context": hybrid_retriever,
        "question": lambda x: x["question"],
        "chat_history": lambda x: x["chat_history"]
    })

    answer_chain = (
        RunnableMap({
            "context": lambda x: format_docs(x["context"]),
            "question": lambda x: x["question"],
            "chat_history": lambda x: format_chat_history(x["chat_history"])
        })
        | prompt
        | llm
        | StrOutputParser()
    )

    final_chain = retrieval_chain.assign(
        answer=answer_chain,
        sources=lambda x: x["context"]
    )

    return final_chain


# STREAMLIT APP
def main():

    st.set_page_config(
        page_title="DocuMind-RAG",
        page_icon="📄"
    )

    st.title("DocuMind-RAG")
    st.write("Multi-Document Retrieval Augmented Generation System")

    # Initialize session states
    if "rag_chain" not in st.session_state:
        st.session_state.rag_chain = None

    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []

    # Display chat history
    for chat in st.session_state.chat_history:
        st.chat_message("user").write(chat["question"])
        st.chat_message("assistant").write(chat["answer"])

    # User Input
    question = st.chat_input("Ask a question about your documents")

    if question and st.session_state.rag_chain:

        with st.spinner("Generating response..."):

            result = st.session_state.rag_chain.invoke({
                "question": question,
                "chat_history": st.session_state.chat_history
            })

            answer = result["answer"]

            # Store in memory
            st.session_state.chat_history.append({
                "question": question,
                "answer": answer
            })

            # Display response
            st.chat_message("user").write(question)
            st.chat_message("assistant").write(answer)

            # Show Sources
            with st.expander("Sources"):
                unique_sources = set()

                for doc in result["sources"]:
                    source_info = f"{doc.metadata['source']} (Page {doc.metadata['page']})"
                    if source_info not in unique_sources:
                        st.write(f"- 📄 {source_info}")
                        unique_sources.add(source_info)

    # SIDEBAR
    with st.sidebar:

        st.header("Model Settings")

        available_models = [
            "gemini-2.5-flash",
            "gemini-2.5-pro",
            "gemini-2.0-flash-001"
        ]

        selected_model = st.selectbox(
            "Select LLM Model",
            available_models
        )

        if "current_model" not in st.session_state:
            st.session_state.current_model = selected_model

        if st.session_state.current_model != selected_model:
            st.session_state.current_model = selected_model
            st.session_state.rag_chain = load_rag_chain(selected_model)

        st.header("Upload Documents")

        uploaded_files = st.file_uploader(
            "Upload Documents (PDF, DOCX, CSV)",
            type=["pdf", "docx", "csv"],
            accept_multiple_files=True
        )

        if st.button("Process Documents"):

            if uploaded_files:

                with st.spinner("Processing Documents..."):

                    documents = []

                    pdf_files = [f for f in uploaded_files if f.name.endswith(".pdf")]
                    docx_files = [f for f in uploaded_files if f.name.endswith(".docx")]
                    csv_files = [f for f in uploaded_files if f.name.endswith(".csv")]

                    if pdf_files:
                        documents.extend(extract_pdf_documents(pdf_files))

                    if docx_files:
                        documents.extend(extract_docx_documents(docx_files))

                    if csv_files:
                        documents.extend(extract_csv_documents(csv_files))

                    chunks = split_documents(documents)
                    create_vector_store(chunks)

                    st.session_state.rag_chain = load_rag_chain(selected_model)

                st.success("Documents processed successfully.")

            else:
                st.warning("Please upload at least one document.")

        # Clear Chat Button
        if st.button("Clear Conversation"):
            st.session_state.chat_history = []
            st.success("Conversation cleared.")


if __name__ == "__main__":
    main()